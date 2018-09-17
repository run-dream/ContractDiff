from __future__ import (absolute_import, division, print_function,
                        unicode_literals)
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import copy
import tools

relative_path = '../../material/2-电子版对照.docx'
path = tools.get_path(relative_path)
document = Document(path)


def table_print(block):
    table = block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text, '  ', end='')
        print("\n")


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


target_path = '../../output/2-diff.docx'
output_path = tools.get_path(target_path)
output_doc = Document()

for block in iter_block_items(document):
    is_paragraph = isinstance(block, Paragraph)
    
    if (is_paragraph):
        print(block.text)
        out_para = output_doc.add_paragraph()
        out_para.style = copy.deepcopy(block.style)
        out_para.paragraph_format = copy.deepcopy(block.paragraph_format)
    else:
        table_print(block)