from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor

import tools
import difflib
import docx_to_txt
import inspect

docx_path = tools.get_path('../../material/2-电子版对照.docx')
docx_txt_path = tools.get_path('../../output/2-电子版对照.txt')
ocr_path = tools.get_path('../../material/2-打印出纸质版.ocr.txt')
result_path = tools.get_path('../../output/2-对照.docx')

ocr_content = ''
with open(ocr_path, 'r') as file:
    ocr_content = file.read()

docx_to_txt.word_to_txt(docx_path, docx_txt_path)

docx_content = ''
with open(docx_txt_path, 'r') as file:
    docx_content = file.read()

differ = difflib.Differ()
diff_content = differ.compare(docx_content, ocr_content)
diff = list(diff_content)


def add_run(paragraph, run, item):
    new_run = paragraph.add_run(item[-1])
    new_run.style = run.style
    font_members = inspect.getmembers(run.font)
    for member in font_members:
        member_name = member[0]
        if (not member_name.startswith('_')):
            value = getattr(run.font, member_name)
            try:
                setattr(new_run.font, member_name, value)
            except Exception:
                pass
    if (item[0] == '-'):
        new_run.font.color.rgb = RGBColor(255, 0, 0)
    elif (item[0] == '+'):
        new_run.font.color.rgb = RGBColor(0, 255, 0)
    elif (item[0] == '?'):
        new_run.font.color.rgb = RGBColor(0, 0, 255)


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


document = Document(docx_path)
document.save(result_path)
document = Document(docx_path)

index = 0
for block in iter_block_items(document):
    is_paragraph = isinstance(block, Paragraph)
    if is_paragraph:
        paragraph = block
        old_runs = paragraph.runs
        if (len(old_runs) == 0):
            continue
        paragraph._p.clear()

        for run in old_runs:
            for char in run.text:
                if index >= len(diff):
                    raise Exception('Out Of Index')
                visited = char == diff[index][-1] and '+' != diff[index][0]
                while not visited and index < len(diff) - 1:
                    add_run(paragraph, run, diff[index])
                    index += 1
                    visited = (char == diff[index][-1]
                               and diff[index][0] != '+')
                add_run(paragraph, run, diff[index])
                index += 1
    else:
        table = block
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    old_runs = paragraph.runs
                    if (len(old_runs) == 0):
                        continue
                    paragraph._p.clear()

                    for run in old_runs:
                        for char in run.text:
                            if index >= len(diff):
                                raise Exception('Out Of Index')
                            visited = (char == diff[index][-1]
                                       and diff[index][0] != '+')
                            while not visited and index < len(diff) - 1:
                                add_run(paragraph, run, diff[index])
                                index += 1
                                visited = (char == diff[index][-1]
                                           and diff[index][0] != '+')
                            add_run(paragraph, run, diff[index])
                            index += 1

document.save(result_path)
print('OK')