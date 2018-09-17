from __future__ import (absolute_import, division, print_function,
                        unicode_literals)
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor
import tools
import difflib
import inspect

source_path = '../../material/2-电子版对照.docx'
source_path = tools.get_path(source_path)

relative_path = '../../output/2-电子版对照.docx'
path = tools.get_path(relative_path)

document = Document(source_path)
document.save(path)
document = Document(path)


def table_print(block):
    table = block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text, '  ', end='')
        print("\n")


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


filepath1 = '../../material/2-电子版对照.txt'
filepath2 = '../../material/2-打印出纸质版.ocr.txt'


def get_diff(path1, path2):
    text1 = tools.get_text(path1)
    text2 = tools.get_text(path2)
    diff = difflib.Differ()
    result_iter = diff.compare(text1, text2)
    return result_iter


result_iter = get_diff(filepath1, filepath2)
diff = list(result_iter)
index = 0
paragraph_no = 0
for block in iter_block_items(document):
    is_paragraph = isinstance(block, Paragraph)
    if (is_paragraph):
        paragraph_no += 1
        text = block.text
        old_runs = block.runs

        print('This is paragragh {} runs{} content{}'.format(
            paragraph_no, len(old_runs), text))

        if (len(old_runs) == 0):
            continue
        old_run = block.runs[0]
        block._p.clear()
        while(diff[index][-1] == '\n'):
            index += 1
        while (diff[index][-1] != '\n'):
            diff_text = ''
            identity = diff[index][0]
            while (diff[index][-1] != '\n' and identity == diff[index][0]):
                diff_text += diff[index][-1]
                index += 1
            new_run = block.add_run(diff_text)
            print(diff_text)
            new_run.style = old_run.style
            font_members = inspect.getmembers(old_run.font)
            for member in font_members:
                member_name = member[0]
                if (not member_name.startswith('_')):
                    value = getattr(old_run.font, member_name)
                    try:
                        setattr(new_run.font, member_name, value)
                    except:
                        pass

            if (identity == '-'):
                new_run.font.color.rgb = RGBColor(255, 0, 0)
            elif (identity == '+'):
                new_run.font.color.rgb = RGBColor(0, 255, 0)
    else:
        table_print(block)

document.save(path)